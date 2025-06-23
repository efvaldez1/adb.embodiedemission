import streamlit as st
import pdfplumber
import pandas as pd
import re
import pytesseract
from pdf2image import convert_from_path
from io import BytesIO
from PIL import Image
import tempfile
import os
import camelot
from pathlib import Path
from sentence_transformers import SentenceTransformer, util
import numpy as np
import plotly.express as px
import matplotlib.pyplot as plt
from docx import Document
from openpyxl import load_workbook
from fpdf import FPDF

# --- Original Utility Functions ---
def convert_df_to_csv(df):
    """
    Convert DataFrame to CSV for download.
    """
    return df.to_csv(index=False).encode('utf-8')

def normalize_string(s):
    """
    Normalize strings by lowercasing, removing punctuation, and extra spaces.
    (Ensuring it's present from your original script's logic)
    """
    s = str(s).lower()
    s = re.sub(r'[^\w\s]', ' ', s)
    s = re.sub(r'\s+', ' ', s)
    return s.strip()

# --- NEW: File Conversion Functions ---

def convert_word_to_pdf(uploaded_file):
    """Converts an uploaded DOCX file to a temporary PDF file."""
    try:
        temp_dir = tempfile.mkdtemp()
        temp_docx_path = os.path.join(temp_dir, uploaded_file.name)
        with open(temp_docx_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        document = Document(temp_docx_path)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=10)

        for para in document.paragraphs:
            pdf.multi_cell(0, 5, para.text)

        for table in document.tables:
            if not table.columns: continue
            col_widths = [int(pdf.w / len(table.columns)) - 2] * len(table.columns)
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    if i < len(col_widths):
                        pdf.cell(col_widths[i], 10, cell.text, border=1)
                pdf.ln()
            pdf.ln()

        temp_pdf_path = os.path.join(temp_dir, f"{os.path.splitext(uploaded_file.name)[0]}.pdf")
        pdf.output(temp_pdf_path)
        
        return temp_pdf_path
    except Exception as e:
        st.error(f"Error converting Word file '{uploaded_file.name}': {e}")
        return None


def convert_excel_to_pdf(uploaded_file):
    """Converts an uploaded XLSX file to a temporary PDF file."""
    try:
        temp_dir = tempfile.mkdtemp()
        temp_xlsx_path = os.path.join(temp_dir, uploaded_file.name)
        with open(temp_xlsx_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        wb = load_workbook(temp_xlsx_path)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=8)

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            pdf.add_page(orientation='L')
            pdf.cell(0, 10, sheet_name, 0, 1, 'C')
            
            if ws.max_column > 0:
                col_widths = [int((pdf.w - 2 * pdf.l_margin) / ws.max_column)] * ws.max_column
                for row in ws.iter_rows():
                    for i, cell in enumerate(row):
                        cell_value = str(cell.value) if cell.value is not None else ""
                        if i < len(col_widths):
                           pdf.cell(col_widths[i], 7, cell_value, border=1)
                    pdf.ln()

        temp_pdf_path = os.path.join(temp_dir, f"{os.path.splitext(uploaded_file.name)[0]}.pdf")
        pdf.output(temp_pdf_path)

        return temp_pdf_path
    except Exception as e:
        st.error(f"Error converting Excel file '{uploaded_file.name}': {e}")
        return None

# --- Original PDF Extraction Logic ---
# (Functions from sbert_final.py are preserved and will be called with UI data)

def extract_country_from_pdf(pdf_path):
    """Extract the country name from the first three pages of the PDF."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages[:3]:
                text = page.extract_text()
                if text:
                    match = re.search(r'(?i)Country:\s*([A-Za-z ]+)', text)
                    if match:
                        return match.group(1).strip()
    except Exception:
        pass
    return "Country Not Found"

def extract_text_from_pdf(pdf_path):
    """Extract text from each page of the PDF using pdfplumber."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = ''
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += f'\n {page_text}'
            return text
    except Exception:
        return ""

def extract_text_with_ocr(pdf_path):
    """Use OCR (pytesseract) on PDF pages converted to images."""
    try:
        images = convert_from_path(pdf_path)
        text = ''
        for image in images:
            text += pytesseract.image_to_string(image)
        return text
    except Exception as e:
        st.warning(f"OCR processing failed for {os.path.basename(pdf_path)}. This might happen if the PDF is not a scanned document or poppler is not installed. Error: {e}")
        return ""

def extract_tables_with_pdfplumber(pdf_path):
    """Extract table data from all pages using pdfplumber with multiple strategies."""
    tables = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_tables = page.extract_tables()
                if not page_tables:
                    page_tables = page.extract_tables(table_settings={"vertical_strategy": "text", "horizontal_strategy": "text"})
                if not page_tables:
                    page_tables = page.extract_tables(table_settings={"vertical_strategy": "lines", "horizontal_strategy": "lines"})
                tables.extend(page_tables)
    except Exception:
        pass
    return tables

def extract_tables_with_camelot(pdf_path):
    """Extract table data from all pages using Camelot."""
    try:
        tables = camelot.read_pdf(pdf_path, pages='all', flavor='lattice')
        if len(tables) == 0 or all(table.df.empty for table in tables):
            tables = camelot.read_pdf(pdf_path, pages='all', flavor='stream')
        return [table.df for table in tables]
    except Exception:
        return []

def parse_quantity(quantity_value):
    """Parse quantity values handling various formats including commas and whitespace."""
    if quantity_value is None: return None
    qty_str = str(quantity_value).strip().replace(',', '')
    if not qty_str: return None
    if isinstance(quantity_value, (int, float)): return float(quantity_value)
    clean_qty = re.sub(r'[^\d.]', '', qty_str)
    try:
        return float(clean_qty) if clean_qty else None
    except ValueError:
        return None

# MODIFIED to accept UI data
def is_target_material(text, target_materials, material_variations):
    """Check if text contains any of our target materials or variations."""
    text_lower = text.lower()
    for material in target_materials:
        for variation in material_variations.get(material, [material]):
            if variation in text_lower:
                return True, material
    return False, None

# MODIFIED to accept UI data
def process_table(table, target_materials, material_variations):
    """Process table data focusing on concrete, steel, and asphalt materials."""
    if not table or len(table) < 2: return None
    header = [str(h).strip() if h is not None else f"Unnamed_{i}" for i, h in enumerate(table[0])]
    df = pd.DataFrame(table[1:], columns=header)
    
    renamed_cols = {}
    for col in df.columns:
        col_lower = str(col).lower()
        if any(term in col_lower for term in ['quantity', 'qty']): renamed_cols[col] = 'Quantity'
        elif any(term in col_lower for term in ['unit', 'uom']): renamed_cols[col] = 'Unit'
        elif any(term in col_lower for term in ['desc', 'material', 'item']): renamed_cols[col] = 'Material'
    df = df.rename(columns=renamed_cols)

    if "Material" not in df.columns or "Quantity" not in df.columns or "Unit" not in df.columns:
        return None

    df = df.dropna(subset=['Material', 'Quantity', 'Unit'])
    df['Quantity'] = df['Quantity'].apply(parse_quantity)
    df = df.dropna(subset=['Quantity'])

    material_info = df['Material'].apply(lambda x: is_target_material(str(x), target_materials, material_variations))
    df['is_target'] = material_info.apply(lambda x: x[0])
    df['Material_Type'] = material_info.apply(lambda x: x[1])
    
    return df[df['is_target']]

# MODIFIED to accept UI data
def extract_material_rows_from_text(text, target_materials, material_variations, material_units):
    """Extract material information from unstructured text."""
    lines = text.split('\n')
    extracted_items = []
    all_units = [unit for units_list in material_units.values() for unit in units_list]
    unit_pattern = r'\b(?:' + '|'.join(re.escape(u) for u in all_units) + r')\b'
    qty_pattern = r'\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\b'

    for line in lines:
        if len(line.strip()) < 5: continue
        is_target, material_type = is_target_material(line, target_materials, material_variations)
        if is_target:
            qty_matches = re.findall(qty_pattern, line)
            unit_matches = re.findall(unit_pattern, line, re.IGNORECASE)
            if qty_matches and unit_matches:
                extracted_items.append({
                    'material': line.strip(),
                    'quantity': parse_quantity(qty_matches[0]),
                    'unit': unit_matches[0],
                    'material_type': material_type
                })
    return pd.DataFrame(extracted_items) if extracted_items else pd.DataFrame()

# MODIFIED to accept UI data
def extract_boq_from_pdf(pdf_path, target_materials, material_variations, material_units):
    """Main BOQ extraction pipeline for a single PDF."""
    extracted_dfs = []

    tables_plumber = extract_tables_with_pdfplumber(pdf_path)
    for table in tables_plumber:
        processed_df = process_table(table, target_materials, material_variations)
        if processed_df is not None and not processed_df.empty: extracted_dfs.append(processed_df)
    
    if not extracted_dfs:
        tables_camelot = extract_tables_with_camelot(pdf_path)
        for table_df in tables_camelot:
            table_list = [table_df.columns.to_list()] + table_df.values.tolist()
            processed_df = process_table(table_list, target_materials, material_variations)
            if processed_df is not None and not processed_df.empty: extracted_dfs.append(processed_df)

    if not extracted_dfs:
        text = extract_text_from_pdf(pdf_path)
        if text:
            text_df = extract_material_rows_from_text(text, target_materials, material_variations, material_units)
            if not text_df.empty: extracted_dfs.append(text_df)

    if not extracted_dfs:
        text_ocr = extract_text_with_ocr(pdf_path)
        if text_ocr:
            ocr_df = extract_material_rows_from_text(text_ocr, target_materials, material_variations, material_units)
            if not ocr_df.empty: extracted_dfs.append(ocr_df)

    if not extracted_dfs:
        return pd.DataFrame()
    
    combined_df = pd.concat(extracted_dfs, ignore_index=True)
    return combined_df

# MODIFIED to accept UI data
def extract_boq_from_files(boq_file, target_materials, material_variations, material_units):
    """Main function to extract BOQ data from uploaded files (PDF, DOCX, XLSX)."""
    temp_pdf_path = None
    file_extension = os.path.splitext(boq_file.name)[1].lower()

    # --- File Conversion Logic ---
    if file_extension == '.pdf':
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(boq_file.getvalue())
            temp_pdf_path = temp_file.name
    elif file_extension == '.docx':
        temp_pdf_path = convert_word_to_pdf(boq_file)
    elif file_extension in ['.xlsx', '.xls']:
        temp_pdf_path = convert_excel_to_pdf(boq_file)
    else:
        st.error(f"Unsupported file format: {boq_file.name}")
        return pd.DataFrame(), None

    if temp_pdf_path is None:
        return pd.DataFrame(), None

    try:
        extracted_df = extract_boq_from_pdf(temp_pdf_path, target_materials, material_variations, material_units)
        country_name = extract_country_from_pdf(temp_pdf_path)
    finally:
        if os.path.exists(temp_pdf_path):
            os.remove(temp_pdf_path)

    # Standardize column names
    col_mapping = {'material': 'Material', 'quantity': 'Quantity', 'unit': 'Unit', 'material_type': 'Material_Type'}
    extracted_df = extracted_df.rename(columns=lambda c: col_mapping.get(c, c))
    
    extracted_df["Source File"] = boq_file.name
    return extracted_df, country_name

# --- Original Emission Factors & GHG Calculation ---
country_to_region = {
    "Kazakhstan": "Asia", "Kyrgyzstan": "Asia", "Tajikistan": "Asia", "Turkmenistan": "Asia",
    "Uzbekistan": "Asia", "China": "Asia", "Democratic People's Republic of Korea": "Asia",
    "Japan": "Asia", "Mongolia": "Asia", "Republic of Korea": "Asia", "Brunei Darussalam": "Asia",
    "Cambodia": "Asia", "Indonesia": "Asia", "Lao People's Democratic Republic": "Asia", 
    "Malaysia": "Asia", "Myanmar": "Asia", "Philippines": "Asia", "Singapore": "Asia",
    "Thailand": "Asia", "Timor-Leste": "Asia", "Viet Nam": "Asia", "Afghanistan": "Asia",
    "Bangladesh": "Asia", "Bhutan": "Asia", "India": "Asia", "Iran": "Asia", "Maldives": "Asia",
    "Nepal": "Asia", "Pakistan": "Asia", "Sri Lanka": "Asia", "Australia": "Oceania",
    "New Zealand": "Oceania", "Fiji": "Oceania", "Papua New Guinea": "Oceania"
}

def extract_country_from_text(text, country_to_region):
    for country in country_to_region.keys():
        if country.lower() in text.lower():
            return country
    return None

def extract_numeric_value(value):
    if pd.isna(value): return None
    match = re.match(r'(\d+\.?\d*)', str(value))
    return float(match.group(1)) if match else None

def extract_uncertainty(gwp_str):
    match = re.search(r'±\s*(\d+\.?\d*)%', str(gwp_str))
    return float(match.group(1)) if match else None

def convert_gwp_to_kg(gwp_str, declared_unit, mass_per_m3=None):
    gwp_match = re.match(r'(\d+\.?\d*)', str(gwp_str))
    if not gwp_match: raise ValueError(f"Invalid GWP format: {gwp_str}")
    gwp = float(gwp_match.group(1))
    unit_lower = str(declared_unit).lower()
    if 'm3' in unit_lower:
        if mass_per_m3 is None or mass_per_m3 == 0: raise ValueError("Mass per 1 m3 is required and cannot be zero for m3 units.")
        return gwp / mass_per_m3
    elif any(u in unit_lower for u in ['t', 'ton', 'tonne']):
        return gwp / 1000
    elif 'kg' in unit_lower:
        quantity_match = re.search(r'(\d+\.?\d*)', unit_lower)
        quantity = float(quantity_match.group(1)) if quantity_match else 1.0
        return gwp / quantity
    else:
        raise ValueError(f"Unsupported declared unit: {declared_unit}")

def convert_boq_to_kg(quantity, unit, material_type=None, density=None):
    unit_upper = unit.strip().upper() if unit else ""
    densities = {'concrete': 2400, 'steel': 7850, 'asphalt': 2300}
    density_val = density if pd.notna(density) else densities.get(str(material_type).lower())

    if not quantity or not unit_upper or not density_val: return None
    if any(u in unit_upper for u in ['T', 'TON', 'TONNE', 'MT']): return quantity * 1000
    if 'KG' in unit_upper: return quantity
    if any(u in unit_upper for u in ['CM', 'M3', 'CUM', 'M³']): return quantity * density_val
    if any(u in unit_upper for u in ['M2', 'SM', 'SQM', 'M²']):
        thickness = 0.05 if material_type == 'asphalt' else 0.15
        return quantity * thickness * density_val
    if any(u in unit_upper for u in ['M', 'LM', 'ML']):
        if material_type == 'steel':
            diameter = 0.016
            return quantity * (np.pi * (diameter/2)**2) * density_val
        else:
            return quantity * (0.1*0.1) * density_val
    return None

def process_emission_factors(df, country_name, country_to_region):
    required_columns = ['Category', 'Declared Unit', 'Mass per 1 m3', 'Average GWP_Global', 'Average GWP_Asia', 'Average GWP_Oceania', 'Density']
    if not all(col in df.columns for col in required_columns):
        st.error(f"Missing required columns in emission factors file: {', '.join(c for c in required_columns if c not in df.columns)}")
        return None, None

    region = country_to_region.get(country_name, "Global")
    gwp_column = f'Average GWP_{region}' if f'Average GWP_{region}' in df.columns and df[f'Average GWP_{region}'].notna().all() else 'Average GWP_Global'
    
    df['Mass per 1 m3'] = df['Mass per 1 m3'].apply(extract_numeric_value)
    df['Density'] = df['Density'].apply(extract_numeric_value)
    df['GWP Uncertainty'] = df[gwp_column].apply(extract_uncertainty)
    
    gwp_kg_values = []
    for _, row in df.iterrows():
        try:
            gwp_val = convert_gwp_to_kg(row[gwp_column], row['Declared Unit'], row['Mass per 1 m3'])
            gwp_kg_values.append(gwp_val)
        except (ValueError, TypeError):
            gwp_kg_values.append(None)
    df['GWP (kgCO2e/kg)'] = gwp_kg_values
    
    return df.rename(columns={'Category': 'EC3 Material Category'}), region

# --- Original Matching & Emissions Calculation (but MODIFIED to accept UI data) ---
model = SentenceTransformer('all-MiniLM-L6-v2')
pd.set_option('display.float_format', '{:.2f}'.format)

def match_materials_with_categories(boq_data, emission_factors, exact_matches_from_ui):
    """Match BOQ materials with EC3 emission factor categories."""
    matched_results = []
    ec3_categories = emission_factors["EC3 Material Category"].tolist()
    if not ec3_categories: return []
    category_embeddings = model.encode(ec3_categories, convert_to_tensor=True)
    
    exact_matches_normalized = {normalize_string(k): v for k, v in exact_matches_from_ui.items()}

    for item in boq_data:
        material_str = str(item.get("Material", ""))
        material_normalized = normalize_string(material_str)
        category, similarity = None, 0.0
        
        if material_normalized in exact_matches_normalized:
            category = exact_matches_normalized[material_normalized]
            similarity = 1.0
        else:
            material_embedding = model.encode(material_str, convert_to_tensor=True)
            similarities = util.cos_sim(material_embedding, category_embeddings)[0].cpu().numpy()
            best_match_index = np.argmax(similarities)
            if best_match_index < len(ec3_categories):
                category = ec3_categories[best_match_index]
                similarity = similarities[best_match_index]

        if category:
            item.update({"Category": category, "Similarity": similarity})
            matched_results.append(item)
    return matched_results

def calculate_ghg_emissions(matched_results, emission_factors, region):
    """Calculate GHG emissions based on matched materials and quantities."""
    emissions_data = []
    for result in matched_results:
        emission_row_df = emission_factors[emission_factors["EC3 Material Category"] == result["Category"]]
        if emission_row_df.empty: continue
        emission_row = emission_row_df.iloc[0]
        
        quantity_kg = convert_boq_to_kg(result["Quantity"], result["Unit"], result.get("Material_Type"), emission_row.get("Density"))
        
        emissions, unsupported_note, gwp_value_per_kg = None, None, None
        if quantity_kg is not None:
            gwp_value_per_kg = emission_row['GWP (kgCO2e/kg)']
            if pd.notna(gwp_value_per_kg):
                emissions = quantity_kg * gwp_value_per_kg
            else:
                unsupported_note = "GWP factor not available"
        else:
            unsupported_note = f"Unsupported unit or missing density: {result['Unit']}"
            
        emissions_data.append({
            "BoQ Material": result["Material"],
            "EC3 Category": result["Category"],
            "Material_Type": result.get("Material_Type"),
            "BoQ Quantity (KG)": quantity_kg,
            "Calculated GHG Emissions (kg CO2e)": emissions,
            "EC3 Regional Average GWP (kgCO2e/kg)": gwp_value_per_kg,
            "Region": region,
            "Similarity": f"{result['Similarity'] * 100:.2f}%",
            "GWP Uncertainty": f"{emission_row['GWP Uncertainty']:.1f}%" if pd.notna(emission_row['GWP Uncertainty']) else "N/A",
            "Source File": result.get("Source File"),
            "Unsupported Unit": unsupported_note
        })
    total_emissions = pd.DataFrame(emissions_data)["Calculated GHG Emissions (kg CO2e)"].sum()
    return emissions_data, total_emissions

# --- Original Dashboard & UI Functions (but MODIFIED) ---
def display_dashboard(emissions_df, all_countries):
    st.markdown("---")
    st.header("📊 Summary Dashboard")
    
    if all_countries:
        st.info(f"Detected Countries: {', '.join(all_countries)}")

    total_emissions = pd.to_numeric(emissions_df["Calculated GHG Emissions (kg CO2e)"], errors='coerce').sum()
    col1, col2 = st.columns(2)
    col1.metric("Total GHG Emissions", f"{total_emissions:,.2f} kg CO2e")

    if 'Material_Type' in emissions_df.columns:
        material_emissions = emissions_df.dropna(subset=['Calculated GHG Emissions (kg CO2e)']).groupby("Material_Type")["Calculated GHG Emissions (kg CO2e)"].sum().reset_index()
        if not material_emissions.empty:
            fig = px.pie(material_emissions, values="Calculated GHG Emissions (kg CO2e)", names="Material_Type", title="Emissions by Material Type")
            col2.plotly_chart(fig, use_container_width=True)

# --- NEW: UI Setup Function to handle editable fields ---
def setup_sidebar_and_ui():
    st.sidebar.header("1. Upload Files")
    emission_file = st.sidebar.file_uploader("Upload Emission Factor Excel File", type=["xlsx"])
    boq_files = st.sidebar.file_uploader("Upload BoQ Files (PDF, DOCX, XLSX)", type=["pdf", "docx", "xlsx"], accept_multiple_files=True)
    
    st.sidebar.header("2. Configure Material Search")

    DEFAULT_TARGET_MATERIALS = ['concrete', 'steel', 'asphalt']
    DEFAULT_MATERIAL_VARIATIONS = {
        'concrete': ['concrete', 'cement', 'rcc', 'pcc', 'precast', 'cast-in-situ', 'reinforced cement', 'c30', 'c20', 'c25', 'c35'],
        'steel': ['steel', 'reinforcement', 'rebar', 'iron bar', 'metal', 'ms rod', 'structural steel', 'ductile iron'],
        'asphalt': ['asphalt', 'bituminous', 'bitumen', 'tack coat', 'prime coat', 'wearing course']
    }
    DEFAULT_MATERIAL_UNITS = {
        'concrete': ['m3', 'm³', 'cum', 'cubic meter', 'cm'],
        'steel': ['kg', 'ton', 'mt', 't', 'tonne', 'nos.'],
        'asphalt': ['m2', 'm²', 'sqm', 'square meter', 'sm', 'm3', 'm³', 'kg']
    }
    DEFAULT_EXACT_MATCHES = {
    "Steel Reinforcement / Iron Bar (1/2‚Äù Round)": "Reinforcing Bar",
    "Steel Reinforcement": "Reinforcing Bar",
    "M. Steel Bar (G-40 & 60)": "Reinforcing Bar",
    "Reinforcement as per AASHTO M-31 Grade 60": "Reinforcing Bar",
    "Reinforcement (Structural Shapes) as per ASTM-A-36": "Reinforcing Bar",
    "50mm - Expansion Joint - Indigenous Type (Steel Plates)": "Plate Steel",
    "Steel bar D12mm (ASTM)": "Reinforcing Bar",
    "Stainless Steel Tank": "Steel Suspension Assemblies",
    "Recovered steel from existing deck slabs": "Merchant Bar (MBQ)",
    "Maintenance of steel and RCC Railing": "Open Web Steel Joists",
    "Metal Beam Crash Barrier": "Composite and Form Decks",
    "Tubular Steel Railing": "Cold Formed Framing",
    "Supply & erection of MS Galvanized octagonal pole": "Hot-Rolled Sections",
    "Asphaltic base course plant mix (Class A)": "Asphalt",
    "Asphaltic Concrete for wearing course (Class A)": "Asphalt",
    "Cut-back asphalt for bituminous prime coat": "Asphalt",
    "Cut-back asphalt for bituminous tack coat": "Asphalt",
    "Scarifying existing bituminous surface": "Asphalt",
    "Dense Graded Bituminous Macadam (DGBM)": "Asphalt",
    "Bituminous Concrete": "Asphalt",
    "Concrete Class A1 (Elevated)": "Ready Mix",
    "Concrete Class A1 (On ground)": "Ready Mix",
    "Concrete Class A1 (Onground)": "Ready Mix",
    "Concrete Class A3 (Elevated)": "Ready Mix",
    "Concrete Class A3 (On ground)": "Ready Mix",
    "Concrete Class A3 (Underground)": "Civil Precast Concrete",
    "Concrete Class B": "Ready Mix",
    "Lean Concrete": "Flowable Concrete Fill",
    "Precast Concrete Class D2 425kg/Sq.m (6,000 psi)": "Structural Precast Concrete",
    "Precast Concrete Class D2 425kg/Sq.m (6,000 psi) (Inlcuding Additional Admixtures super plasticizer Sikament 520 (ABS) or equivalent 1.25% by weight of Cement w\c ratio must not exceed 0.32-0.35 and slump should not be more than 160 mm the cost which is deemed to be included in the cost of Concrete)": "Structural Precast Concrete",
    "Plum Concrete (Cyclopean/Rubble)": "Civil Precast Concrete",
    "Reinforced Concrete grade 25 Mpa for bottom slab": "Structural Precast Concrete",
    "Concrete ring for slurry pit √ò100cm": "Utility & Underground Precast Concrete",
    "Concrete ring for mixing tank √ò60cmx0.5m": "Utility & Underground Precast Concrete",
    "Reinforced Cement Concrete Crash Barrier": "Structural Precast Concrete",
    "Cast in Situ Cement Concrete M 20 Kerb": "Structural Precast Concrete",
    "Providing and Laying Reinforced Cement Concrete Pipe NP3": "Utility & Underground Precast Concrete",
    "Rust Removal of Exposed Corroded Rebars (Chemrite Descaler A-28 or Equivalent)": "Metal Surface Treatments",
    "Sealing of Exposed Concrete Gaps with Epoxy Mortar (Chmedur 31 or Equivalent)": "Concrete Additives",
    "SBR Latex modified concrete": "Concrete Additives",
    "Cost difference between SR cement and OPC": "Portland Cement",
    "Additional Admixtures Chermite 520 (ABS) or equivalent": "Concrete Additives",
    "Additional Admixtures Silica Fumes 6% by weight of cement": "Concrete Additives",
    "Plum Concrete (Cyclopean/Rubble) (2:1 concrete stone Ratio)": "Civil Precast Concrete",
    "Concrete Class A3 Onground (from concrete mix plant)": "Ready Mix",
    "Concrete Class A3 Elevated (from concrete mix plant)": "Ready Mix",
    "40% cost of salvaged steel from the deck slab, Barrier and Approach Slab": "Merchant Bar (MBQ)",
    "Dismantling and disposal of structures and obstruction": "Not Applicable",
    "50mm - Expansion Joint - Indigonous Type 15MM thick two vertical steel plates welded 10MM thick round steel bars local (Pakistan Make),PSQCA Certified (As Specified in Drawings)": "Plate Steel"
}

    target_materials_str = st.sidebar.text_area("Target Materials", ', '.join(DEFAULT_TARGET_MATERIALS))
    target_materials = [m.strip().lower() for m in target_materials_str.split(',') if m.strip()]

    material_variations, material_units = {}, {}
    with st.sidebar.expander("Edit Material Variations and Units"):
        for material in target_materials:
            st.markdown(f"**{material.capitalize()}**")
            variations_str = st.text_input(f"Variations for {material}", ', '.join(DEFAULT_MATERIAL_VARIATIONS.get(material, [])), key=f"vars_{material}")
            material_variations[material] = [v.strip().lower() for v in variations_str.split(',') if v.strip()]
            units_str = st.text_input(f"Units for {material}", ', '.join(DEFAULT_MATERIAL_UNITS.get(material, [])), key=f"units_{material}")
            material_units[material] = [u.strip().lower() for u in units_str.split(',') if u.strip()]

    st.header("3. Review and Edit Exact Material Matches")
    match_df = pd.DataFrame(DEFAULT_EXACT_MATCHES.items(), columns=["BoQ Material Description", "EC3 Material Category"])
    
    with st.expander("Edit Exact Matches"):
        edited_match_df = st.data_editor(match_df, num_rows="dynamic", use_container_width=True)
        edited_match_df.dropna(inplace=True)
        exact_matches = dict(zip(edited_match_df["BoQ Material Description"], edited_match_df["EC3 Material Category"]))

    return emission_file, boq_files, target_materials, material_variations, material_units, exact_matches

# --- Main Application (Adapted from original sbert_final.py)---
def main():
    st.set_page_config(layout="wide")
    st.title("🏗️ Embodied Emissions Calculator for Construction")
    st.markdown("This tool analyzes Bill of Quantity (BoQ) documents to estimate the embodied carbon footprint of construction materials.")

    emission_file, boq_files, target_materials, material_variations, material_units, exact_matches = setup_sidebar_and_ui()

    if st.button("Calculate Emissions", type="primary"):
        if emission_file and boq_files:
            with st.spinner("Processing files..."):
                try:
                    df_emission = pd.read_excel(emission_file)
                except Exception as e:
                    st.error(f"Failed to read emission factor file: {e}")
                    return

                all_boq_data, all_countries = [], set()
                
                for boq_file in boq_files:
                    st.info(f"Processing {boq_file.name}...")
                    # Pass UI-defined material settings to the extraction function
                    extracted_df, country_name = extract_boq_from_files(boq_file, target_materials, material_variations, material_units)
                    
                    if country_name and country_name != "Country Not Found":
                        all_countries.add(country_name)
                    
                    if not extracted_df.empty:
                        all_boq_data.append(extracted_df)
                    else:
                        st.warning(f"No target materials found in {boq_file.name}.")

                if all_boq_data:
                    combined_boq_df = pd.concat(all_boq_data, ignore_index=True)
                    st.success(f"Extracted {len(combined_boq_df)} relevant material items from all files.")
                    
                    country_for_region = next(iter(all_countries), None)
                    processed_emission_df, region = process_emission_factors(df_emission, country_for_region, country_to_region)
                    
                    if processed_emission_df is not None:
                        boq_records = combined_boq_df.to_dict('records')
                        # Pass UI-defined exact matches to the matching function
                        matched_results = match_materials_with_categories(boq_records, processed_emission_df, exact_matches)
                        emissions_data, total_emissions = calculate_ghg_emissions(matched_results, processed_emission_df, region)
                        
                        if emissions_data:
                            emissions_df = pd.DataFrame(emissions_data)
                            display_dashboard(emissions_df, all_countries)
                            
                            st.header("📄 Detailed Emissions Report")
                            st.dataframe(emissions_df)
                            
                            csv = convert_df_to_csv(emissions_df)
                            st.download_button("Download Report as CSV", csv, "emissions_report.csv", "text/csv")
                        else:
                            st.error("Could not calculate emissions. Check if BoQ items match the emission factor categories.")
                else:
                    st.error("No relevant material data could be extracted from the uploaded files.")
        else:
            st.warning("Please upload an emission factors file and at least one BoQ file.")

if __name__ == "__main__":
    main()
